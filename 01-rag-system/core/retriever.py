from __future__ import annotations

from io import BytesIO
from pathlib import Path

import streamlit as st
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from core.chunking import split_documents
from core.utils import file_signature, format_context_sections, list_base_knowledge_files, preview_text, source_label, weak_retrieval
from core.vector_store import VectorIndex, build_vector_index


EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"


def load_base_documents() -> list[Document]:
    documents = []
    for path in list_base_knowledge_files():
        text = path.read_text(encoding="utf-8")
        documents.append(Document(page_content=text, metadata={"source": path.name, "type": "base"}))
    return documents


@st.cache_resource(show_spinner=False)
def get_base_index() -> VectorIndex | None:
    base_docs = load_base_documents()
    return build_vector_index(split_documents(base_docs), EMBED_MODEL, origin="base")


def _parse_pdf(uploaded_file) -> list[Document]:
    reader = PdfReader(BytesIO(uploaded_file.getvalue()))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(
                Document(
                    page_content=text,
                    metadata={"source": uploaded_file.name, "page": page_number, "type": "upload"},
                )
            )
    return pages


def _parse_docx(uploaded_file) -> list[Document]:
    doc = DocxDocument(BytesIO(uploaded_file.getvalue()))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    if not paragraphs:
        return []
    return [Document(page_content="\n".join(paragraphs), metadata={"source": uploaded_file.name, "type": "upload"})]


def _parse_txt(uploaded_file) -> list[Document]:
    text = uploaded_file.getvalue().decode("utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": uploaded_file.name, "type": "upload"})]


def parse_uploaded_documents(uploaded_files: list | None) -> list[Document]:
    documents: list[Document] = []
    for uploaded_file in uploaded_files or []:
        suffix = Path(uploaded_file.name).suffix.lower()
        try:
            if suffix == ".pdf":
                documents.extend(_parse_pdf(uploaded_file))
            elif suffix == ".docx":
                documents.extend(_parse_docx(uploaded_file))
            elif suffix == ".txt":
                documents.extend(_parse_txt(uploaded_file))
        except Exception:
            st.warning(f"Skipped {uploaded_file.name} because it could not be processed cleanly.")
    return documents


def build_uploaded_index(uploaded_files: list | None) -> tuple[VectorIndex | None, int, int]:
    uploaded_docs = parse_uploaded_documents(uploaded_files)
    if not uploaded_docs:
        return None, 0, 0
    chunks = split_documents(uploaded_docs)
    index = build_vector_index(chunks, EMBED_MODEL, origin="uploaded")
    return index, len(uploaded_docs), len(chunks)


def update_uploaded_index_state(uploaded_files: list | None) -> None:
    current_signature = file_signature(uploaded_files)
    if st.session_state.get("upload_signature", "") == current_signature:
        return
    if uploaded_files:
        with st.spinner("Indexing uploaded files..."):
            index, doc_count, chunk_count = build_uploaded_index(uploaded_files)
        st.session_state.upload_index = index
        st.session_state.upload_doc_count = doc_count
        st.session_state.upload_chunk_count = chunk_count
    else:
        st.session_state.upload_index = None
        st.session_state.upload_doc_count = 0
        st.session_state.upload_chunk_count = 0
    st.session_state.upload_signature = current_signature


def _search_index(index: VectorIndex | None, query: str, limit: int) -> list[Document]:
    if not index:
        return []
    return index.vectorstore.similarity_search(query, k=limit)


def retrieve_shared_context(question: str, base_index: VectorIndex | None, uploaded_index: VectorIndex | None = None) -> dict:
    per_index_limit = 3 if len(question.split()) > 10 else 2
    candidates = []
    seen = set()
    for index in filter(None, [uploaded_index, base_index]):
        for doc in _search_index(index, question, per_index_limit):
            key = (doc.metadata.get("source"), doc.metadata.get("page"), doc.page_content[:160])
            if key in seen:
                continue
            seen.add(key)
            candidates.append(doc)
    candidates = candidates[:5]
    return {
        "documents": candidates,
        "context": format_context_sections(candidates, limit=460),
        "retrieved_chunks": len(candidates),
        "sources": list(dict.fromkeys(source_label(doc.metadata) for doc in candidates)),
        "source_cards": [{"label": source_label(doc.metadata), "preview": preview_text(doc.page_content, 260)} for doc in candidates],
        "weak_retrieval": weak_retrieval(question, candidates),
    }
