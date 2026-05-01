from __future__ import annotations

import base64
import html
import io
from datetime import datetime
from statistics import mean

import streamlit as st
import streamlit.components.v1 as components

from features.voice_output import render_voice_output

STARTER_PROMPTS = [
    "English: What are the main KYC requirements for banks?",
    "\u0c24\u0c46\u0c32\u0c41\u0c17\u0c41: \u0c2c\u0c4d\u0c2f\u0c3e\u0c02\u0c15\u0c41\u0c32\u0c4d\u0c32\u0c4b KYC \u0c15\u0c4b\u0c38\u0c02 \u0c05\u0c35\u0c38\u0c30\u0c2e\u0c48\u0c28 \u0c2a\u0c4d\u0c30\u0c27\u0c3e\u0c28 \u0c2a\u0c24\u0c4d\u0c30\u0c3e\u0c32\u0c41 \u0c0f\u0c2e\u0c3f\u0c1f\u0c3f?",
    "\u4e2d\u6587: \u94f6\u884cKYC\u5408\u89c4\u6700\u91cd\u8981\u7684\u8981\u6c42\u662f\u4ec0\u4e48\uff1f",
    "Espa\u00f1ol: \u00bfCu\u00e1les son los requisitos principales de KYC para bancos?",
    "Fran\u00e7ais: Quelles sont les principales exigences KYC pour les banques ?",
    "\u0420\u0443\u0441\u0441\u043a\u0438\u0439: \u041a\u0430\u043a\u043e\u0432\u044b \u043e\u0441\u043d\u043e\u0432\u043d\u044b\u0435 \u0442\u0440\u0435\u0431\u043e\u0432\u0430\u043d\u0438\u044f KYC \u0434\u043b\u044f \u0431\u0430\u043d\u043a\u043e\u0432?",
]

TECH_ITEMS = [
    ("\U0001F916", "LLM", "OpenAI / Fine-Tuned Model"),
    ("\U0001F9E0", "Embeddings", "text-embedding-3"),
    ("\U0001F5C2", "Vector DB", "FAISS / Hybrid"),
    ("\U0001F517", "Framework", "LangChain"),
    ("\U0001F5A5", "UI", "Streamlit"),
    ("\u2699", "Infra", "Python / FastAPI"),
]


def _as_html_text(text: str) -> str:
    return html.escape(str(text)).replace("\n", "<br>")


def _greeting() -> tuple[str, str]:
    hour = datetime.now().hour
    if 5 <= hour < 12:
        return "Good morning", chr(0x1F324)
    if 12 <= hour < 17:
        return "Good afternoon", chr(0x2600)
    if 17 <= hour < 21:
        return "Good evening", chr(0x1F306)
    return "Good night", chr(0x1F319)


def inject_premium_css() -> None:
    st.markdown(
        """
<style>
:root{--bg:#F6F8FC;--bg-soft:#EEF3FA;--text:#0B1220;--border:rgba(15,23,42,.08);--navy:#123A6F}
html, body, [data-testid="stAppViewContainer"]{background:var(--bg)!important;color:var(--text)!important}
.stApp{background:linear-gradient(180deg,#F8FBFF 0%,#F4F7FC 100%)!important}
.block-container{max-width:none!important;padding:.1rem .16rem 8.2rem!important}
[data-testid="stHeader"]{background:transparent!important}
#MainMenu, footer{visibility:hidden}
[data-testid="stSidebar"]{background:var(--bg-soft)!important;border-right:1px solid var(--border)!important;min-width:300px!important}
.sidebar-brand{display:flex;gap:12px;align-items:center;margin:4px 0 14px}
.sidebar-title{font-size:17px;font-weight:900;line-height:1.1;color:#08245A}
.sidebar-caption{font-size:11px;color:#5E78A8;margin-top:4px;font-weight:600}
.side-search{height:42px;border:1px solid var(--border);border-radius:14px;background:#FFF;color:#94A3B8;display:flex;align-items:center;padding:0 12px;margin:8px 0 18px;font-size:13px}
.sidebar-section-label{margin:2px 0 10px;font-size:11px;font-weight:900;letter-spacing:.08em;text-transform:uppercase;color:#7A8EAA}
[data-testid="stSidebar"] .stButton > button,.stButton > button[id*="starter-"]{width:100%!important;border-radius:12px!important;border:1px solid #1D4ED8!important;background:linear-gradient(135deg,#2563EB,#0EA5E9)!important;color:#FFF!important;font-weight:800!important}
.product-shell{background:rgba(255,255,255,.86);border:1px solid var(--border);border-radius:22px;box-shadow:0 16px 34px rgba(15,23,42,.06);padding:18px;margin:0 0 8px}
.greeting-row{display:flex;align-items:center;gap:8px;margin:0 0 10px 2px;color:#08245A;font-weight:850}
.greeting-pill{width:30px;height:30px;border-radius:10px;background:#FFF7ED;border:1px solid #FED7AA;display:flex;align-items:center;justify-content:center}
.greeting-sub{font-weight:600;color:#29456F;font-size:15px;margin-bottom:18px}
.hero-card{display:grid;grid-template-columns:112px 1fr 280px;gap:20px;align-items:center;min-height:142px;border:1px solid rgba(37,99,235,.10);border-radius:22px;background:linear-gradient(135deg,rgba(255,255,255,.97),rgba(241,247,255,.90));padding:18px 24px}
.hero-title{font-size:30px;font-weight:900;letter-spacing:-.04em;color:#08245A;margin:0 0 8px}
.hero-copy{font-size:15px;color:#274871;line-height:1.55}
.proof-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin:16px 0}
.proof-card,.info-card{background:#FFF;border:1px solid var(--border);border-radius:22px;padding:18px}
.proof-icon{width:42px;height:42px;border-radius:14px;background:#EFF6FF;color:#2563EB;display:flex;align-items:center;justify-content:center;font-size:21px;margin-bottom:10px}
.proof-title,.info-title{font-weight:900;color:#123A6F;font-size:15px}
.proof-text,.info-copy{font-size:12.5px;color:#64748B;line-height:1.55;margin-top:6px}
.product-info-grid{display:grid;grid-template-columns:1.15fr 1fr;gap:14px;margin:14px 0 22px}
.tech-row{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;margin-top:12px}
.tech-item{text-align:center;font-size:11px;color:#34516F}
.tech-ico{width:34px;height:34px;border-radius:12px;background:#EFF6FF;margin:0 auto 6px;display:flex;align-items:center;justify-content:center;color:#2563EB;font-size:17px}
.user-row{display:flex;justify-content:flex-end;margin:0 0 14px}
.user-bubble{max-width:72%;background:var(--navy);color:#FFF;border-radius:16px 16px 4px 16px;padding:12px 18px;font-weight:700;line-height:1.5;font-size:14px}
.ai-wrap{display:grid;grid-template-columns:54px 1fr;gap:14px;align-items:start;margin:0 0 12px}
.ai-globe{font-size:44px;line-height:1}
.answer-shell,.thinking-shell{background:#FFF;border:1px solid var(--border);border-radius:18px;padding:18px 20px;box-shadow:0 10px 24px rgba(15,23,42,.045);font-size:14.5px;line-height:1.72}
.thinking-shell{display:flex;align-items:center;gap:10px;min-height:72px}
.thinking-dot{width:10px;height:10px;border-radius:50%;background:#2563EB}
.thinking-text{color:#274871;font-weight:700}
.meta-pills{display:flex;gap:8px;flex-wrap:wrap;margin-bottom:12px}
.meta-pill{font-size:11px;font-weight:900;border-radius:999px;padding:5px 10px;border:1px solid var(--border);background:#F8FBFF;color:#123A6F}
.meta-pill.green{background:#ECFDF5;color:#047857;border-color:#BBF7D0}
.starter-label{margin:6px 0 8px;color:#123A6F;font-size:13px;font-weight:900}
.composer-shell-static{position:sticky!important;bottom:8px!important;z-index:70!important;background:#FFF!important;border:1px solid rgba(37,99,235,.14)!important;border-radius:16px!important;box-shadow:0 12px 30px rgba(15,23,42,.08)!important;padding:8px 10px!important;margin:8px 0 0!important}
.composer-shell form{border:none!important;background:transparent!important}
.composer-marker{display:none!important}
.composer-pending{visibility:hidden!important;opacity:0!important;pointer-events:none!important}
.composer-ready{visibility:visible!important;opacity:1!important;pointer-events:auto!important}
.composer-row [data-testid="column"]{display:flex;align-items:center}
.composer-row [data-testid="column"] > div{width:100%}
.composer-shell div[data-testid="stPopover"] button,
.composer-shell div[data-testid="stSelectbox"] > div[data-baseweb="select"] > div{min-height:44px!important;border-radius:14px!important;background:#FFF!important;border:1px solid rgba(15,23,42,.08)!important;color:#123A6F!important}
.composer-shell [data-testid="stTextInput"] > div > div{background:#FFF!important;border:1px solid rgba(37,99,235,.14)!important;border-radius:14px!important}
.composer-shell [data-testid="stTextInput"] input{min-height:44px!important;padding:10px 14px!important}
.composer-shell [data-testid="stFormSubmitButton"] button{min-height:44px!important;border-radius:14px!important;background:#123A6F!important;color:#FFF!important;border:none!important}
@media (max-width:1100px){.hero-card{grid-template-columns:80px 1fr}.proof-grid,.product-info-grid,.tech-row{grid-template-columns:1fr}}
</style>
        """,
        unsafe_allow_html=True,
    )


def render_header() -> None:
    greeting, icon = _greeting()
    st.markdown(
        f"""
        <div class="product-shell">
          <div class="greeting-row"><span class="greeting-pill" data-greeting-icon>{icon}</span><span data-greeting-text>{html.escape(greeting)}</span></div>
          <div class="greeting-sub">I&#39;m your Banking &amp; Finance Copilot.</div>
          <div class="hero-card">
            <div style="font-size:78px;">&#127758;</div>
            <div>
              <div class="hero-title">Banking &amp; Finance AI Copilot</div>
              <div class="hero-copy">Your intelligent assistant for banking, compliance, and financial intelligence.<br>
              Ask anything related to AML, KYC, FDIC, Basel III, RBI, regulations, and more. Upload PDF, DOCX, TXT, and image files for grounded answers from your reference material.</div>
            </div>
            <div style="font-size:82px; text-align:center;">&#127974;&#128737;&#65039;</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_homepage_product_intro() -> None:
    proof_cards = [
        ("&#128737;&#65039;", "Source-grounded answers", "Every response is backed by trusted documents with citations."),
        ("&#127760;", "Multilingual workflows", "Use multilingual text, PDF, DOCX, TXT, image, and voice workflows across the app."),
        ("&#9989;", "Evaluation-ready AI system", "Confidence scores, latency, and retrieval insights for transparency."),
    ]
    proof_html = "".join(
        f'<div class="proof-card"><div class="proof-icon">{icon}</div><div class="proof-title">{title}</div><div class="proof-text">{copy}</div></div>'
        for icon, title, copy in proof_cards
    )
    tech_html = "".join(
        f'<div class="tech-item"><div class="tech-ico">{icon}</div><strong>{html.escape(label)}</strong><br>{html.escape(value)}</div>'
        for icon, label, value in TECH_ITEMS
    )
    st.markdown(
        f"""
        <div class="product-shell">
          <div class="proof-grid">{proof_html}</div>
          <div class="product-info-grid">
            <div class="info-card">
              <div class="info-title">About this Copilot</div>
              <div class="info-copy">A product-grade AI assistant for banking and finance professionals. Ask questions, upload documents, and review grounded answers with confidence signals and transparent sources.</div>
            </div>
            <div class="info-card">
              <div class="info-title">Tech Stack</div>
              <div class="tech-row">{tech_html}</div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_welcome_card() -> None:
    render_homepage_product_intro()


def render_starter_prompts() -> str | None:
    selected = None
    st.markdown('<div class="starter-label">Popular prompts</div>', unsafe_allow_html=True)
    cols = st.columns(3)
    for index, prompt in enumerate(STARTER_PROMPTS):
        with cols[index % 3]:
            if st.button(prompt, key=f"starter-{index}", use_container_width=True):
                selected = prompt
    return selected


def render_about_section() -> None:
    return None


def render_stack_section() -> None:
    return None


def render_sidebar_brand() -> None:
    st.markdown(
        """
        <div class="sidebar-brand">
          <div style="font-size:40px;">&#127758;</div>
          <div>
            <div class="sidebar-title">Banking &amp; Finance<br>AI Copilot</div>
            <div class="sidebar-caption">Grounded. Trusted. Intelligent.</div>
          </div>
        </div>
        <div class="side-search">&#8989; Search conversations...</div>
        """,
        unsafe_allow_html=True,
    )


def render_sidebar_summary(base_doc_count: int, upload_doc_count: int, upload_chunk_count: int) -> None:
    st.markdown(
        f"""
        <div style="font-size:12px;color:#48618A;padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Base files</span><strong style="color:#123A6F;">{base_doc_count}</strong></div>
        <div style="font-size:12px;color:#48618A;padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Uploaded docs</span><strong style="color:#123A6F;">{upload_doc_count}</strong></div>
        <div style="font-size:12px;color:#48618A;padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Uploaded chunks</span><strong style="color:#123A6F;">{upload_chunk_count}</strong></div>
        """,
        unsafe_allow_html=True,
    )


def render_session_insights(messages: list[dict]) -> None:
    assistant_messages = [m for m in messages if m.get("role") == "assistant"]
    avg_latency = round(mean(m.get("latency_ms", 0) for m in assistant_messages)) if assistant_messages else 0
    avg_chunks = round(mean(m.get("retrieved_chunks", 0) for m in assistant_messages), 1) if assistant_messages else 0
    st.markdown(
        f"""
        <div style="font-size:12px;color:#48618A; padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Answers</span><strong style="color:#123A6F;">{len(assistant_messages)}</strong></div>
        <div style="font-size:12px;color:#48618A; padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Avg latency</span><strong style="color:#123A6F;">{avg_latency} ms</strong></div>
        <div style="font-size:12px;color:#48618A; padding:4px 0;display:grid;grid-template-columns:1fr auto;gap:8px;"><span>Avg chunks</span><strong style="color:#123A6F;">{avg_chunks}</strong></div>
        """,
        unsafe_allow_html=True,
    )


def render_user_message(text: str) -> None:
    st.markdown(f'<div class="user-row"><div class="user-bubble">{_as_html_text(text)}</div></div>', unsafe_allow_html=True)


def _copy_button(answer: str) -> None:
    encoded = base64.b64encode(answer.encode("utf-8")).decode("utf-8")
    components.html(
        f"""
        <button onclick="(function(btn){{
            const raw = atob('{encoded}');
            const bytes = Uint8Array.from(raw, c => c.charCodeAt(0));
            const text = new TextDecoder('utf-8').decode(bytes);
            navigator.clipboard.writeText(text).then(() => {{ btn.innerHTML='Copied'; setTimeout(() => btn.innerHTML='Copy', 1200); }});
        }})(this)"
        style="width:100%;height:38px;border:1px solid rgba(15,23,42,.10);border-radius:12px;background:#fff;color:#123A6F;font-weight:700;cursor:pointer;">
          Copy
        </button>
        """,
        height=42,
    )


def _docx_bytes(answer: str, source_cards: list[dict]) -> bytes:
    try:
        from docx import Document
    except Exception:
        return b""
    doc = Document()
    doc.add_heading("Banking & Finance Copilot Answer", level=1)
    doc.add_paragraph(answer)
    if source_cards:
        doc.add_heading("Sources", level=2)
        for card in source_cards[:5]:
            doc.add_paragraph(f"{card.get('label', 'Source')}: {str(card.get('preview', ''))[:300]}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_assistant_message(message: dict, message_key: str, simplified_answers: bool, show_source_cards: bool, show_auto_comparison: bool) -> None:
    answer = str(message.get("answer", "")).strip() if simplified_answers else str(message.get("answer", ""))
    backend = str(message.get("backend", "OpenAI")).replace("GPT-4o", "OpenAI")
    st.markdown('<div class="ai-wrap"><div class="ai-globe">&#127758;</div><div>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="answer-shell">
          <div class="meta-pills">
            <span class="meta-pill">{html.escape(backend)}</span>
            <span class="meta-pill">{html.escape(str(message.get("latency_ms", "--")))} ms</span>
            <span class="meta-pill">{html.escape(str(message.get("retrieved_chunks", "--")))} chunks</span>
            <span class="meta-pill green">{html.escape(str(message.get("confidence", "High")))} confidence</span>
          </div>
          {_as_html_text(answer)}
        </div>
        """,
        unsafe_allow_html=True,
    )

    cols = st.columns([1.15, 1.0, 1.0, 1.0, 0.72, 0.72, 4.0])
    with cols[0]:
        render_voice_output(answer, message_key, lang_hint=str(message.get("voice_lang_hint", "") or ""))
    with cols[1]:
        _copy_button(answer)
    with cols[2]:
        st.download_button("TXT", data=answer.encode("utf-8"), file_name=f"copilot-answer-{message_key}.txt", mime="text/plain", key=f"dl-txt-{message_key}", use_container_width=True)
    with cols[3]:
        docx_bytes = _docx_bytes(answer, message.get("source_cards", []))
        if docx_bytes:
            st.download_button("DOCX", data=docx_bytes, file_name=f"copilot-answer-{message_key}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl-docx-{message_key}", use_container_width=True)
        else:
            st.button("DOCX", key=f"dl-docx-unavail-{message_key}", disabled=True, use_container_width=True)
    with cols[4]:
        st.button("\U0001F44D", key=f"up-{message_key}", use_container_width=True)
    with cols[5]:
        st.button("\U0001F44E", key=f"down-{message_key}", use_container_width=True)
    st.markdown('<div style="font-size:11px;color:#64748B;margin:8px 0 0 2px;">AI can make mistakes. Verify important information with official sources.</div>', unsafe_allow_html=True)
    st.markdown("</div></div>", unsafe_allow_html=True)


def render_assistant_thinking() -> None:
    st.markdown(
        """
        <div class="ai-wrap">
          <div class="ai-globe">&#127758;</div>
          <div class="thinking-shell"><span class="thinking-dot"></span><span class="thinking-text">Thinking...</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_footer() -> None:
    return None


def enforce_composer_pin() -> None:
    return None
